"""
Policy routes — Sprint 6 chunks 6.2, 6.11, 6.14, 6.15.
========================================================
POST  /chat/resume              — resume an interrupted HITL graph
GET   /api/policies             — list user's policy drafts
GET   /api/policies/{id}        — get a single policy draft
PATCH /api/policies/{id}        — update policy content (edit persistence)
GET   /api/policies/{id}/download — download as .md or .docx
POST  /api/policies/{id}/finalize — enqueue DOCX generation job

Refs: PLAN.md chunks 6.2, 6.11, 6.14, 6.15; ADR-0007; US-011-015.
"""

from __future__ import annotations

import asyncio
import io
import logging
import os
import re
import uuid
from datetime import UTC, datetime
from typing import Annotated, Literal

from fastapi import APIRouter, Depends, HTTPException, Query, Request, status
from fastapi.responses import StreamingResponse
from opentelemetry import trace
from pydantic import BaseModel, ConfigDict, Field
from slowapi import Limiter
from slowapi.util import get_remote_address

from apps.api.auth.clerk import ClerkUser, verify_clerk_token
from apps.api.db import AppDbPool, AppDbPoolDep

logger = logging.getLogger(__name__)
tracer = trace.get_tracer(__name__)
router = APIRouter(tags=["policies"])

_policies_limiter = Limiter(key_func=get_remote_address, default_limits=[])


def _policies_rate_limit() -> str:
    return os.environ.get("POLICIES_RATE_LIMIT", "60/minute")


# ── Schemas ──────────────────────────────────────────────────────────────────


class ResumeRequest(BaseModel):
    """POST /chat/resume body shape (Sprint 6 chunk 6.2)."""

    model_config = ConfigDict(extra="forbid")

    thread_id: str = Field(max_length=128, description="The thread_id of the interrupted graph.")
    decision: Literal["approve", "edit", "reject"]
    edited_content: str | None = Field(default=None, max_length=100_000)
    rejection_reason: str | None = Field(default=None, max_length=2000)


class PolicyOut(BaseModel):
    """One policy draft row."""

    model_config = ConfigDict(extra="forbid")

    id: str
    user_id: str
    policy_type: str
    title: str
    content: str
    version: int
    finalized: bool
    thread_id: str | None = None
    created_at: str
    updated_at: str


class PoliciesListOut(BaseModel):
    model_config = ConfigDict(extra="forbid")
    policies: list[PolicyOut] = Field(default_factory=list)
    count: int = 0


class PolicyPatch(BaseModel):
    """PATCH /api/policies/{id} body — edit persistence (chunk 6.11)."""

    model_config = ConfigDict(extra="forbid")

    content: str = Field(max_length=100_000)
    title: str | None = Field(default=None, max_length=200)


# ── DB helpers ───────────────────────────────────────────────────────────────


async def _upsert_policy_draft(
    pool: AppDbPool,
    *,
    user_id: str,
    draft_id: str,
    policy_type: str,
    title: str,
    content: str,
    version: int,
    finalized: bool,
    thread_id: str | None,
) -> None:
    """Insert or update a policy draft in the policy_drafts table."""

    async with pool.connection() as conn:
        async with conn.transaction():
            await conn.execute(
                "SELECT set_config('app.current_user_id', %s, true)",
                (user_id,),
            )
            await conn.execute(
                """
                INSERT INTO policy_drafts
                    (id, user_id, policy_type, title, content, version, finalized, thread_id)
                VALUES
                    (%s::uuid, %s, %s, %s, %s, %s, %s, %s)
                ON CONFLICT (id) DO UPDATE SET
                    title      = EXCLUDED.title,
                    content    = EXCLUDED.content,
                    version    = EXCLUDED.version,
                    finalized  = EXCLUDED.finalized,
                    updated_at = now()
                """,
                (draft_id, user_id, policy_type, title, content, version, finalized, thread_id),
            )


async def _insert_revision(
    pool: AppDbPool,
    *,
    user_id: str,
    policy_id: str,
    content: str,
    version: int,
    source: str,
) -> None:
    """Insert a row into policy_revisions for the internal revision log (chunk 6.11)."""

    async with pool.connection() as conn:
        async with conn.transaction():
            await conn.execute(
                "SELECT set_config('app.current_user_id', %s, true)",
                (user_id,),
            )
            await conn.execute(
                """
                INSERT INTO policy_revisions
                    (id, policy_id, user_id, content, version, source)
                VALUES
                    (%s::uuid, %s::uuid, %s, %s, %s, %s)
                """,
                (str(uuid.uuid4()), policy_id, user_id, content, version, source),
            )


async def _fetch_policy(pool: AppDbPool, *, user_id: str, policy_id: str) -> PolicyOut | None:
    async with pool.connection() as conn:
        async with conn.transaction():
            await conn.execute(
                "SELECT set_config('app.current_user_id', %s, true)",
                (user_id,),
            )
            async with conn.cursor() as cur:
                await cur.execute(
                    """
                    SELECT id::text, user_id, policy_type, title, content,
                           version, finalized, thread_id, created_at, updated_at
                    FROM policy_drafts
                    WHERE user_id = %s AND id = %s::uuid
                    """,
                    (user_id, policy_id),
                )
                row = await cur.fetchone()
                if row is None:
                    return None
                return _row_to_policy(row)


async def _list_policies(pool: AppDbPool, *, user_id: str) -> list[PolicyOut]:
    async with pool.connection() as conn:
        async with conn.transaction():
            await conn.execute(
                "SELECT set_config('app.current_user_id', %s, true)",
                (user_id,),
            )
            async with conn.cursor() as cur:
                await cur.execute(
                    """
                    SELECT id::text, user_id, policy_type, title, '',
                           version, finalized, thread_id, created_at, updated_at
                    FROM policy_drafts
                    WHERE user_id = %s
                    ORDER BY updated_at DESC
                    LIMIT 100
                    """,
                    (user_id,),
                )
                rows = await cur.fetchall()
                return [_row_to_policy(r) for r in rows]


async def _update_policy_content(
    pool: AppDbPool, *, user_id: str, policy_id: str, content: str, title: str | None
) -> PolicyOut | None:
    async with pool.connection() as conn:
        async with conn.transaction():
            await conn.execute(
                "SELECT set_config('app.current_user_id', %s, true)",
                (user_id,),
            )
            async with conn.cursor() as cur:
                if title is not None:
                    await cur.execute(
                        """
                        UPDATE policy_drafts
                        SET content = %s, title = %s, version = version + 1, updated_at = now()
                        WHERE user_id = %s AND id = %s::uuid
                        RETURNING id::text, user_id, policy_type, title, content,
                                  version, finalized, thread_id, created_at, updated_at
                        """,
                        (content, title, user_id, policy_id),
                    )
                else:
                    await cur.execute(
                        """
                        UPDATE policy_drafts
                        SET content = %s, version = version + 1, updated_at = now()
                        WHERE user_id = %s AND id = %s::uuid
                        RETURNING id::text, user_id, policy_type, title, content,
                                  version, finalized, thread_id, created_at, updated_at
                        """,
                        (content, user_id, policy_id),
                    )
                row = await cur.fetchone()
                if row is None:
                    return None
                return _row_to_policy(row)


def _row_to_policy(row: tuple) -> PolicyOut:
    (id_, user_id, policy_type, title, content,
     version, finalized, thread_id, created_at, updated_at) = row
    return PolicyOut(
        id=id_,
        user_id=user_id,
        policy_type=policy_type,
        title=title or "",
        content=content or "",
        version=version,
        finalized=bool(finalized),
        thread_id=thread_id,
        created_at=_iso(created_at),
        updated_at=_iso(updated_at),
    )


def _iso(dt: datetime) -> str:
    if dt and dt.tzinfo is None:
        dt = dt.replace(tzinfo=UTC)
    return dt.isoformat() if dt else ""


# ── Routes ───────────────────────────────────────────────────────────────────


@router.get("/api/policies", response_model=PoliciesListOut)
@_policies_limiter.limit(_policies_rate_limit)
async def list_policies(
    request: Request,
    pool: AppDbPoolDep,
    user: Annotated[ClerkUser, Depends(verify_clerk_token)],
) -> PoliciesListOut:
    with tracer.start_as_current_span("policies.list") as span:
        span.set_attribute("user.id", user.user_id)
        policies = await _list_policies(pool, user_id=user.user_id)
        return PoliciesListOut(policies=policies, count=len(policies))


@router.get("/api/policies/{policy_id}", response_model=PolicyOut)
@_policies_limiter.limit(_policies_rate_limit)
async def get_policy(
    request: Request,
    policy_id: str,
    pool: AppDbPoolDep,
    user: Annotated[ClerkUser, Depends(verify_clerk_token)],
) -> PolicyOut:
    with tracer.start_as_current_span("policies.get") as span:
        span.set_attribute("user.id", user.user_id)
        span.set_attribute("policy.id", policy_id)
        policy = await _fetch_policy(pool, user_id=user.user_id, policy_id=policy_id)
        if policy is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Policy not found")
        return policy


@router.patch("/api/policies/{policy_id}", response_model=PolicyOut)
@_policies_limiter.limit(_policies_rate_limit)
async def patch_policy(
    request: Request,
    policy_id: str,
    body: PolicyPatch,
    pool: AppDbPoolDep,
    user: Annotated[ClerkUser, Depends(verify_clerk_token)],
) -> PolicyOut:
    """Edit persistence — save user edits (chunk 6.11, US-012)."""

    with tracer.start_as_current_span("policies.patch") as span:
        span.set_attribute("user.id", user.user_id)
        span.set_attribute("policy.id", policy_id)
        updated = await _update_policy_content(
            pool,
            user_id=user.user_id,
            policy_id=policy_id,
            content=body.content,
            title=body.title,
        )
        if updated is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Policy not found")

        # Record revision for the internal change log.
        try:
            await _insert_revision(
                pool,
                user_id=user.user_id,
                policy_id=policy_id,
                content=body.content,
                version=updated.version,
                source="user_edit",
            )
        except Exception:  # noqa: BLE001
            logger.warning("policy_revision.insert_failed policy_id=%s", policy_id)

        return updated


@router.get("/api/policies/{policy_id}/download")
@_policies_limiter.limit(_policies_rate_limit)
async def download_policy(
    request: Request,
    policy_id: str,
    pool: AppDbPoolDep,
    user: Annotated[ClerkUser, Depends(verify_clerk_token)],
    format: Annotated[
        Literal["md", "docx"],
        Query(description="Download format: 'md' or 'docx'."),
    ] = "md",
) -> StreamingResponse:
    """Download a policy as Markdown or DOCX (Sprint 6 chunk 6.15).

    DOCX conversion uses python-docx. The file is generated on-the-fly
    and streamed — no R2 upload required for v1.
    """

    with tracer.start_as_current_span("policies.download") as span:
        span.set_attribute("user.id", user.user_id)
        span.set_attribute("policy.id", policy_id)
        span.set_attribute("format", format)

        policy = await _fetch_policy(pool, user_id=user.user_id, policy_id=policy_id)
        if policy is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Policy not found")

        safe_title = re.sub(r'[^\w\-]', '_', policy.title)[:50] or "policy"

        if format == "md":
            return StreamingResponse(
                io.BytesIO(policy.content.encode("utf-8")),
                media_type="text/markdown; charset=utf-8",
                headers={
                    "Content-Disposition": (
                        f'attachment; filename="{safe_title}'
                        f'_v{policy.version}.md"'
                    ),
                },
            )

        # DOCX generation (Sprint 6 chunk 6.14). Runs in a thread to
        # avoid blocking the event loop — python-docx does synchronous
        # XML serialisation.
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_501_NOT_IMPLEMENTED,
                detail="python-docx is not installed on this server.",
            )

        def _build_docx() -> bytes:
            doc = Document()
            doc.add_heading(policy.title, level=0)
            doc.add_paragraph(
                f"Policy type: {policy.policy_type} | Version: {policy.version} | "
                f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}"
            )
            doc.add_paragraph("")
            for line in policy.content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    doc.add_paragraph(stripped)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        docx_bytes = await asyncio.to_thread(_build_docx)

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": (
                    f'attachment; filename="{safe_title}'
                    f'_v{policy.version}.docx"'
                ),
            },
        )


__all__ = [
    "PolicyOut",
    "PolicyPatch",
    "ResumeRequest",
    "router",
]
